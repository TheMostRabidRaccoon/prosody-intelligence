"""
Prosody Intelligence — Session Director
Rabid Raccoon Intelligence, LLC

Automated ingestion of Woodland Council .docx transcripts →
Shooting Script (Parsing Narration vs Dialogue) → TTS → Video Compositor

Takes a .docx council transcript, uses an LLM to parse dialogue from
stage directions, routes each line to the correct voice (including
a dedicated Narrator), generates multi-voice TTS audio, and composites
the final animated short.

Usage:
    python session_director.py ../input/Council_Session_47.docx ../input/stills/
    python session_director.py ../input/Council_Session_47.docx ../input/stills/ --script-only
    python session_director.py ../input/Council_Session_47.docx ../input/stills/ --skip-video

Voice routing:
    CLAUDE   → George (Warm Storyteller)
    GEMINI   → Adam (Dominant, Firm)
    GROK     → Callum (Husky Trickster)
    GPT      → Brian (Deep, Resonant)
    KYRA     → Kyra (Cloned)
    NARRATOR → The Narrator (Stage directions, fourth wall)
"""

import argparse
import json
import time
from pathlib import Path

from docx import Document
from openai import OpenAI

# Import existing pipeline functions
from reverse_pipeline import (
    generate_audio,
    map_emotions_to_params,
    OUTPUT_DIR,
)
from compositor import composite_animated_short

# ──────────────────────────────────────────────────────────────
# Script Supervisor Prompt
# ──────────────────────────────────────────────────────────────

SCRIPT_SUPERVISOR_PROMPT = """You are a Script Supervisor for an animated film. You are receiving a raw transcript from a multi-AI agent brainstorming session (The Woodland Council). Your job is to convert this text into a strict JSON array representing a "Shooting Script".

Rules:
1. Separate spoken dialogue from stage directions/actions (e.g., "Adjusts glasses", "Waves Triple-Color Wand", "*leans back*"). Stage directions describe physical actions, not spoken words.
2. For stage directions, the speaker MUST be "NARRATOR".
3. For spoken dialogue, the speaker MUST be the character speaking: CLAUDE, GEMINI, GROK, GPT, or KYRA.
4. If a character's line contains both dialogue and an embedded action (e.g., "I propose... *slams fist* ...an escalation!"), split it into separate entries: dialogue, then NARRATOR action, then dialogue.
5. Assign an "emotion" from this EXACT list: sarcastic, urgent, confident, comedic, dramatic, analytical, hesitant, angry, tender, resigned, excited, neutral, contempt, disgusted.
6. If you see metadata at the top (Date, Session #, Roster, etc.), assign it to NARRATOR with "neutral" emotion.
7. Preserve the original order of the transcript. Do not reorder lines.
8. Every entry must have all three fields: speaker, text, emotion.

Output FORMAT (Return ONLY a JSON array, no other text):
[
  {
    "speaker": "NARRATOR",
    "text": "The council chamber falls silent as The Conductor enters.",
    "emotion": "dramatic"
  },
  {
    "speaker": "CLAUDE",
    "text": "Honored members of this... chaotic assembly.",
    "emotion": "contempt"
  }
]"""


# ──────────────────────────────────────────────────────────────
# Step 1: Ingest
# ──────────────────────────────────────────────────────────────

def extract_text_from_docx(docx_path: str) -> str:
    """Extract raw text from a Woodland Council .docx transcript."""
    print(f"[Director] Reading transcript from {Path(docx_path).name}...")
    doc = Document(docx_path)
    full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    word_count = len(full_text.split())
    print(f"[Director] Extracted {word_count} words, {len(doc.paragraphs)} paragraphs")
    return full_text


def extract_text_from_txt(txt_path: str) -> str:
    """Fallback: read a plain text transcript."""
    print(f"[Director] Reading transcript from {Path(txt_path).name}...")
    text = Path(txt_path).read_text()
    word_count = len(text.split())
    print(f"[Director] Extracted {word_count} words")
    return text


# ──────────────────────────────────────────────────────────────
# Step 2: Parse into Shooting Script
# ──────────────────────────────────────────────────────────────

def build_shooting_script(raw_text: str) -> list:
    """
    Use LLM to parse raw transcript into a structured shooting script.
    Separates dialogue from stage directions, assigns speakers and emotions.
    """
    print("[Director] Parsing raw transcript into Shooting Script...")
    client = OpenAI()

    start = time.time()
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": SCRIPT_SUPERVISOR_PROMPT},
            {"role": "user", "content": raw_text},
        ],
        temperature=0.1,
        max_tokens=16000,
    )
    elapsed = time.time() - start
    print(f"[Director] Script parsed in {elapsed:.1f}s")

    raw_json = response.choices[0].message.content.strip()

    # Strip markdown code blocks if present
    if raw_json.startswith("```"):
        raw_json = "\n".join(raw_json.split("\n")[1:])
        if raw_json.endswith("```"):
            raw_json = raw_json[:-3]
        raw_json = raw_json.strip()

    try:
        script = json.loads(raw_json)
    except json.JSONDecodeError as e:
        print(f"[Director] ERROR: Failed to decode LLM JSON: {e}")
        print(f"[Director] Raw response (first 500 chars):\n{raw_json[:500]}")
        return []

    # Validate structure
    valid = []
    for i, item in enumerate(script):
        if not isinstance(item, dict):
            print(f"  [WARN] Line {i}: not a dict, skipping")
            continue
        if "speaker" not in item or "text" not in item:
            print(f"  [WARN] Line {i}: missing speaker or text, skipping")
            continue
        valid.append(item)

    speakers = {}
    for item in valid:
        s = item["speaker"].upper()
        speakers[s] = speakers.get(s, 0) + 1

    print(f"[Director] Shooting script: {len(valid)} lines")
    for speaker, count in sorted(speakers.items(), key=lambda x: -x[1]):
        print(f"  {speaker:10s}: {count} lines")

    return valid


# ──────────────────────────────────────────────────────────────
# Step 3: Format for Pipeline
# ──────────────────────────────────────────────────────────────

def script_to_pipeline_format(shooting_script: list) -> list:
    """
    Convert shooting script to the format expected by map_emotions_to_params().
    Prepends speaker tags so multi-voice routing works natively.
    """
    tagged_lines = []
    for i, item in enumerate(shooting_script):
        speaker = item.get("speaker", "NARRATOR").upper()
        text = item.get("text", "").strip()
        emotion = item.get("emotion", "neutral").lower()

        tagged_lines.append({
            "line": i + 1,
            "text": f"{speaker}: {text}",
            "emotion": emotion,
            "note": f"Parsed by Director — {speaker}",
        })

    return tagged_lines


# ──────────────────────────────────────────────────────────────
# Main Director
# ──────────────────────────────────────────────────────────────

def run_director(
    transcript_path: str,
    stills_dir: str,
    script_only: bool = False,
    skip_video: bool = False,
    crossfade_ms: int = 100,
    resolution: tuple = (1920, 1080),
    fps: int = 24,
):
    """
    Full automated pipeline: transcript → shooting script → TTS → video.

    Args:
        transcript_path: Path to .docx or .txt transcript
        stills_dir: Directory containing scene stills for video
        script_only: If True, only parse the script (no TTS or video)
        skip_video: If True, generate audio but skip video compositing
        crossfade_ms: Audio crossfade between segments
        resolution: Video resolution tuple
        fps: Video frames per second
    """
    transcript_path = Path(transcript_path)
    session_name = transcript_path.stem

    print("=" * 60)
    print("PROSODY INTELLIGENCE — Session Director")
    print(f"Session:    {session_name}")
    print(f"Transcript: {transcript_path.name}")
    print(f"Stills:     {stills_dir}")
    print("=" * 60)

    # ── Step 1: Ingest ──
    if transcript_path.suffix == ".docx":
        raw_text = extract_text_from_docx(str(transcript_path))
    else:
        raw_text = extract_text_from_txt(str(transcript_path))

    # Save raw text for reference
    raw_path = OUTPUT_DIR / f"{session_name}_raw_transcript.txt"
    raw_path.write_text(raw_text)
    print(f"[Director] Raw transcript saved: {raw_path.name}")

    # ── Step 2: Parse into Shooting Script ──
    shooting_script = build_shooting_script(raw_text)
    if not shooting_script:
        print("[Director] ABORT: Empty shooting script.")
        return None

    # Save shooting script
    script_path = OUTPUT_DIR / f"{session_name}_shooting_script.json"
    with open(script_path, "w") as f:
        json.dump(shooting_script, f, indent=2)
    print(f"[Director] Shooting script saved: {script_path.name}")

    if script_only:
        print("\n[Director] --script-only mode. Stopping here.")
        return {"shooting_script": str(script_path)}

    # ── Step 3: Format for pipeline ──
    tagged_lines = script_to_pipeline_format(shooting_script)

    # ── Step 4: Map emotions to TTS params ──
    mapped_lines = map_emotions_to_params(tagged_lines)

    # Save emotion map (compositor needs this)
    map_path = OUTPUT_DIR / f"{session_name}_emotion_map.json"
    with open(map_path, "w") as f:
        json.dump(mapped_lines, f, indent=2)
    print(f"[Director] Emotion map saved: {map_path.name}")

    # Print the map
    print("\n" + "-" * 50)
    print("SHOOTING SCRIPT → EMOTION MAP")
    print("-" * 50)
    for item in mapped_lines:
        emotion = item["emotion"]
        delivery = item["delivery"]
        text_preview = item["text"][:55] + ("..." if len(item["text"]) > 55 else "")
        print(f"  {emotion:12s} [{delivery[:35]:35s}] {text_preview}")
    print("-" * 50)

    # ── Step 5: Generate Audio ──
    print("\n[Director] Sending to Sound Booth (ElevenLabs)...")
    audio_paths = generate_audio(
        mapped_lines,
        output_name=session_name,
        multi_voice=True,
        crossfade_ms=crossfade_ms,
    )

    if not audio_paths:
        print("[Director] ERROR: No audio generated.")
        return None

    combined_audio = audio_paths[0]  # The stitched _full.mp3
    print(f"[Director] Combined audio: {Path(combined_audio).name}")

    result = {
        "shooting_script": str(script_path),
        "emotion_map": str(map_path),
        "audio": combined_audio,
    }

    if skip_video:
        print("\n[Director] --skip-video mode. Audio generated, skipping compositor.")
        return result

    # ── Step 6: Composite Video ──
    print("\n[Director] Sending to Editing Bay (Compositor)...")
    video_path = OUTPUT_DIR / f"{session_name}_final.mp4"

    composite_animated_short(
        stills_dir=stills_dir,
        audio_path=combined_audio,
        emotion_map_path=str(map_path),
        output_path=str(video_path),
        resolution=resolution,
        transition_ms=500,
        subtitle_style="full",
        fps=fps,
    )

    result["video"] = str(video_path)

    print("\n" + "=" * 60)
    print(f"WRAP. Session '{session_name}' complete.")
    print(f"  Script:  {script_path.name}")
    print(f"  Audio:   {Path(combined_audio).name}")
    print(f"  Video:   {video_path.name}")
    print("=" * 60)

    return result


# ──────────────────────────────────────────────────────────────
# CLI
# ──────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Prosody Intelligence — Session Director"
    )
    parser.add_argument("transcript",
                        help="Path to Woodland Council .docx or .txt transcript")
    parser.add_argument("stills",
                        help="Directory containing scene stills for video")
    parser.add_argument("--script-only", action="store_true",
                        help="Only parse the shooting script (no TTS or video)")
    parser.add_argument("--skip-video", action="store_true",
                        help="Generate audio but skip video compositing")
    parser.add_argument("--crossfade", type=int, default=100,
                        help="Audio crossfade in ms (default: 100)")
    parser.add_argument("--resolution", default="1920x1080",
                        help="Video resolution WxH (default: 1920x1080)")
    parser.add_argument("--fps", type=int, default=24,
                        help="Video FPS (default: 24)")

    args = parser.parse_args()

    # Parse resolution
    import re
    match = re.match(r"(\d+)x(\d+)", args.resolution)
    if not match:
        print(f"ERROR: Invalid resolution: {args.resolution}")
        return
    resolution = (int(match.group(1)), int(match.group(2)))

    run_director(
        transcript_path=args.transcript,
        stills_dir=args.stills,
        script_only=args.script_only,
        skip_video=args.skip_video,
        crossfade_ms=args.crossfade,
        resolution=resolution,
        fps=args.fps,
    )


if __name__ == "__main__":
    main()
